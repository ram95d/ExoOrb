import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Inches
import base64
from mpl_toolkits.mplot3d import Axes3D
import streamlit as st
# Set page configuration
st.set_page_config(page_title="EVs Anticancer Therapeutic Potential Comparison and Metrics", layout="wide")

# Main title
st.title("🔬 EVs Anticancer Therapeutic Potential Comparison and Metrics")

# --- Session State Initialization ---
if "data_df" not in st.session_state:
    st.session_state.data_df = None
if "initial_maximize_selection_guess" not in st.session_state:
    st.session_state.initial_maximize_selection_guess = []
if "last_input_method" not in st.session_state:
    st.session_state.last_input_method = None 

# Sidebar for input data
st.sidebar.header("Input Options")

# Choose between manual entry and file upload
input_method = st.sidebar.radio("Select Input EVs", ["Manual Entry", "Upload Excel File"], key="input_method_radio")

if st.session_state.last_input_method != input_method:
    st.session_state.data_df = None
    st.session_state.initial_maximize_selection_guess = []
    st.session_state.last_input_method = input_method

manual_factor_names = []
method_names = []

if input_method == "Upload Excel File":
    st.sidebar.subheader("Upload Data")
    uploaded_file = st.sidebar.file_uploader("Choose an Excel file", type=["xlsx", "xls"], key="file_uploader")
    
    if uploaded_file is not None:
        try:
            loaded_df = pd.read_excel(uploaded_file, index_col=0)
            st.session_state.data_df = loaded_df 
            
            st.subheader("Uploaded Data")
            st.write(st.session_state.data_df)
            
            method_names = list(st.session_state.data_df.index)
            uploaded_factors = list(st.session_state.data_df.columns)
            
            maximize_keywords = ["yield", "purity", "recovery", "antioxidant", "stability", "rna"]
            st.session_state.initial_maximize_selection_guess = [
                factor for factor in uploaded_factors 
                if any(keyword.lower() in factor.lower() for keyword in maximize_keywords)
            ]
        except Exception as e:
            st.error(f"Error reading the file: {e}")
            st.info("Please make sure your Excel file is properly formatted with EVs in the first column and Parameters in the header row.")
            st.session_state.data_df = None 
            st.session_state.initial_maximize_selection_guess = []
    else:
        st.session_state.data_df = None
        st.session_state.initial_maximize_selection_guess = []

else:
    st.sidebar.subheader("Define EVs and Characterization Parameters")
    num_methods = st.sidebar.number_input("Number of EVs", min_value=2, max_value=20, value=3, key="num_methods")
    num_factors = st.sidebar.number_input("Number of Parameters", min_value=1, max_value=20, value=3, key="num_factors")
    
    st.sidebar.subheader("EVs Names")
    temp_method_names = [] 
    for i in range(num_methods):
        method_name = st.sidebar.text_input(f"EVs {i+1} Name", value=f"EVs {i+1}", key=f"manual_method_name_{i}")
        temp_method_names.append(method_name)
    method_names = temp_method_names 
    
    st.sidebar.subheader("Parameter Names")
    temp_manual_factor_names = []
    for i in range(num_factors):
        factor_name = st.sidebar.text_input(f"Parameters {i+1} Name", value=f"Parameters {i+1}", key=f"manual_factor_name_{i}")
        temp_manual_factor_names.append(factor_name)
    manual_factor_names = temp_manual_factor_names
    
    st.sidebar.subheader("Specify Parameters Optimization Goal")
    if manual_factor_names:
        valid_defaults_for_manual_select = [
            f for f in st.session_state.get('initial_maximize_selection_guess', [])
            if f in manual_factor_names
        ]
        if not valid_defaults_for_manual_select and manual_factor_names:
            default_for_manual_select = manual_factor_names[:len(manual_factor_names)//2]
        else:
            default_for_manual_select = valid_defaults_for_manual_select

        current_manual_maximize_selection = st.sidebar.multiselect(
            "Select Parameters to Maximize",
            manual_factor_names,
            default=default_for_manual_select,
            key="manual_setup_maximize_factors_select"
        )
        st.session_state.initial_maximize_selection_guess = current_manual_maximize_selection
        
        derived_minimize_factors = [f for f in manual_factor_names if f not in st.session_state.initial_maximize_selection_guess]
        st.sidebar.write("Parameters to Minimize (derived automatically):")
        st.sidebar.write(", ".join(derived_minimize_factors) if derived_minimize_factors else "None")
    else:
        st.sidebar.write("Define Parameter names first to classify them.")
        st.session_state.initial_maximize_selection_guess = []

    st.subheader("Enter Parameters Values for Each EVs")
    method_data = {}
    if method_names and manual_factor_names:
        with st.form("data_entry_form"):
            cols = st.columns(len(manual_factor_names))
            for i, factor in enumerate(manual_factor_names):
                cols[i].write(f"**{factor}**")
            
            for method in method_names:
                method_data[method] = {}
                row_cols = st.columns(len(manual_factor_names))
                for i, factor in enumerate(manual_factor_names):
                    method_data[method][factor] = row_cols[i].number_input(
                        f"{method} - {factor}", value=0.0, format="%.6f", key=f"input_{method}_{factor}"
                    )
            submitted = st.form_submit_button("Submit Data")
        
        if submitted:
            if method_data:
                local_df = pd.DataFrame.from_dict(method_data, orient='index')
                if not local_df.empty and manual_factor_names:
                    local_df.columns = manual_factor_names
                st.session_state.data_df = local_df
                st.subheader("Manually Entered Data")
                st.write(st.session_state.data_df)
            else:
                st.warning("No data entered in the form.")
                st.session_state.data_df = None
    else:
        st.info("Please define EVs and parameter names to enter data.")


if st.session_state.data_df is not None and not st.session_state.data_df.empty:
    actual_factors_in_df = list(st.session_state.data_df.columns)
    actual_method_names = list(st.session_state.data_df.index)

    st.sidebar.subheader("Confirm Parameters Classification")
    
    valid_initial_guess_for_confirmation = [
        f for f in st.session_state.get('initial_maximize_selection_guess', []) 
        if f in actual_factors_in_df
    ]
    maximize_factors_selected = st.sidebar.multiselect(
        "Select Parameters to Maximize", 
        actual_factors_in_df,
        default=valid_initial_guess_for_confirmation,
        key="confirm_maximize_select"
    )
    minimize_factors_selected = [factor for factor in actual_factors_in_df if factor not in maximize_factors_selected]
    st.sidebar.write("Parameters to Minimize:")
    st.sidebar.write(", ".join(minimize_factors_selected) if minimize_factors_selected else "None")
    
    st.sidebar.subheader("Weights for Each Parameters")
    weights = {}
    if not actual_factors_in_df:
        st.sidebar.warning("No Parameters available to assign weights.")
    for factor in actual_factors_in_df:
        weights[factor] = st.sidebar.slider(f"{factor} Weight", 0.0, 1.0, 0.10, 0.01, key=f"weight_{factor}")
    w_series = pd.Series(weights)
    w_sum = w_series.sum()
    
    if w_sum == 0:
        st.error("All weights are zero. Please adjust sliders.")
        st.stop()
    
    w_series = w_series / w_sum 
    weights = w_series.to_dict()  
    if st.button("Compare and Rank EVs", key="compare_button"):
        if not actual_factors_in_df:
            st.error("No Parameters defined for comparison.")
        elif not weights:
            st.error("Weights not set for factors.")
        elif not maximize_factors_selected and not minimize_factors_selected:
            st.error("No Parameters classified for maximization or minimization.")
        else:
            df_processed = st.session_state.data_df.copy()
            
            for factor in actual_factors_in_df:
                if 'zeta' in factor.lower() or 'potential' in factor.lower():
                    st.info(f"Detected '{factor}' as potential zeta/potential measurement - will apply special normalization.")
                    continue
                if factor in df_processed.columns and df_processed[factor].min() < 0:
                    shift_value = abs(df_processed[factor].min()) + 1
                    st.info(f"Shifted '{factor}' values by +{shift_value:.2f} to make all values positive before normalization.")
                    df_processed[factor] = df_processed[factor] + shift_value
            
            for factor in maximize_factors_selected:
                if factor in df_processed.columns:
                    if df_processed[factor].max() != 0:
                        df_processed[factor + '_normalized'] = df_processed[factor] / df_processed[factor].max()
                    else: 
                        df_processed[factor + '_normalized'] = 0.0 if not df_processed[factor].empty else pd.Series([0.0]*len(df_processed))


            for factor in minimize_factors_selected:
                if factor in df_processed.columns:
                    series_to_normalize = df_processed[factor]
                    norm_col_name = factor + '_normalized'

                    if 'zeta' in factor.lower() or 'potential' in factor.lower() or (series_to_normalize < 0).any():
                        st.info(f"Applying special normalization for '{factor}' (minimize, may contain negatives).")

                        if (series_to_normalize <= 0).all() and series_to_normalize.min() < 0: 
                             df_processed[norm_col_name] = series_to_normalize / series_to_normalize.min()
                        else: 
                            temp_series = series_to_normalize.replace(0, 1e-9)
                            if temp_series.min() > 0 :
                                df_processed[norm_col_name] = temp_series.min() / temp_series
                            elif temp_series.abs().max() > 0 :
                                df_processed[norm_col_name] = 1 - (temp_series.abs() / temp_series.abs().max())
                            else:
                                df_processed[norm_col_name] = 1.0
                    elif (series_to_normalize >= 0).all():
                        min_val = series_to_normalize.min()
                        max_val = series_to_normalize.max()
                        if max_val == min_val :
                            df_processed[norm_col_name] = 1.0 if min_val == 0 else 0.0
                        elif max_val > 0 :
                             if min_val > 0:
                                 df_processed[norm_col_name] = min_val / series_to_normalize.replace(0, 1e-9)
                             else: 
                                 df_processed[norm_col_name] = 1 - (series_to_normalize / max_val)

                        else: 
                            df_processed[norm_col_name] = 1.0
                    else: 
                        st.warning(f"Parameter '{factor}' has mixed signs and is minimized. Normalizing based on absolute distance to zero.")
                        abs_max = series_to_normalize.abs().max()
                        if abs_max > 0:
                            df_processed[norm_col_name] = 1 - (series_to_normalize.abs() / abs_max)
                        else:
                            df_processed[norm_col_name] = 1.0
            
            normalized_factors_present = [f + '_normalized' for f in actual_factors_in_df if (f + '_normalized') in df_processed.columns]
            
            if not normalized_factors_present:
                st.error("No Parameters were normalized. Please check Parameter classifications and values.")
            else:
                df_processed['Score'] = 0.0 
                for factor in actual_factors_in_df:
                    norm_col = factor + '_normalized'
                    if norm_col in df_processed.columns and factor in weights:
                        df_processed[norm_col] = df_processed[norm_col].fillna(0)
                        df_processed['Score'] += df_processed[norm_col] * weights[factor]
                
                df_processed['Rank'] = df_processed['Score'].rank(ascending=False, method='min')
                
                st.subheader("EVs Rankings")
                ranking_df = df_processed[['Score', 'Rank']].sort_values('Rank')
                st.write(ranking_df)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.subheader("Bar Plot of Scores")
                    fig_bar, ax_bar = plt.subplots(figsize=(10, 6))
                    sorted_df_plot = df_processed.sort_values('Score', ascending=False)
                    ax_bar.bar(sorted_df_plot.index, sorted_df_plot['Score'], color='skyblue')
                    ax_bar.set_ylabel('Score')
                    ax_bar.set_title('EVs Scores (Higher is Better)')
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()
                    st.pyplot(fig_bar)
                
                with col2:
                    st.subheader("Heatmap of Normalized Parameters")
                    if normalized_factors_present and not df_processed[normalized_factors_present].empty:
                        fig_heat, ax_heat = plt.subplots(figsize=(12, 8))
                        sns.heatmap(df_processed[normalized_factors_present], annot=True, cmap='YlGnBu', ax=ax_heat, fmt=".2f")
                        plt.title('Normalized Parameter Values (Higher is Better)')
                        plt.xticks(rotation=45, ha='right')
                        plt.yticks(rotation=0)
                        plt.tight_layout()
                        st.pyplot(fig_heat)
                    else:
                        st.write("No normalized data to display in heatmap.")

                if not df_processed.empty and 'Score' in df_processed.columns and normalized_factors_present:
                    st.subheader("Top 3 EVs Radar Chart")
                    top_n = min(3, len(df_processed))
                    top_methods_df = df_processed.nlargest(top_n, 'Score')
                    
                    if not top_methods_df.empty and not top_methods_df[normalized_factors_present].empty:
                        radar_data = top_methods_df[normalized_factors_present]
                        radar_labels = [factor.replace('_normalized', '') for factor in normalized_factors_present]
                        angles = np.linspace(0, 2 * np.pi, len(radar_labels), endpoint=False).tolist()
                        angles += angles[:1] 

                        fig_radar, ax_radar = plt.subplots(figsize=(10, 8), subplot_kw=dict(projection='polar'))
                        for method_idx, method_name_radar in enumerate(top_methods_df.index):
                            values = radar_data.loc[method_name_radar].fillna(0).tolist()
                            values += values[:1] 
                            ax_radar.plot(angles, values, 'o-', linewidth=2, label=method_name_radar)
                            ax_radar.fill(angles, values, alpha=0.25)
                        
                        ax_radar.set_xticks(angles[:-1])
                        ax_radar.set_xticklabels(radar_labels, fontsize=8)
                        ax_radar.set_yticks(np.arange(0, 1.1, 0.2))
                        ax_radar.set_yticklabels([f"{i:.1f}" for i in np.arange(0, 1.1, 0.2)])
                        plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
                        plt.title('Top EVs Comparison (Normalized Values)')
                        st.pyplot(fig_radar)


                if len(normalized_factors_present) >= 2 and 'Score' in df_processed.columns and 'Rank' in df_processed.columns:
                    st.subheader("3D Scatter Plot: Distance to Ideal Solutions & Scores")
                    
                    if all(col in df_processed.columns for col in normalized_factors_present):
                        ideal_best = df_processed[normalized_factors_present].max()
                        ideal_worst = df_processed[normalized_factors_present].min()
                        
                        df_processed['Distance to Ideal Best'] = np.sqrt(
                            np.sum((df_processed[normalized_factors_present].fillna(0) - ideal_best.fillna(0)) ** 2, axis=1))
                        df_processed['Distance to Ideal Worst'] = np.sqrt(
                            np.sum((df_processed[normalized_factors_present].fillna(0) - ideal_worst.fillna(0)) ** 2, axis=1))

                        cmap = plt.get_cmap('viridis')
                        fig_3d = plt.figure(figsize=(12, 8))
                        ax_3d = fig_3d.add_subplot(111, projection='3d')
                        scatter = ax_3d.scatter(
                            df_processed['Distance to Ideal Best'],
                            df_processed['Distance to Ideal Worst'],
                            df_processed['Score'],
                            c=df_processed['Rank'],
                            cmap=cmap,
                            s=100,
                        )
                
                        for method_name in df_processed.index:
                            x = df_processed.loc[method_name, 'Distance to Ideal Best']
                            y = df_processed.loc[method_name, 'Distance to Ideal Worst']
                            z = df_processed.loc[method_name, 'Score']
                            rank = int(df_processed.loc[method_name, 'Rank'])
                            ax_3d.text(x, y, z, str(rank), size=7, color='white', ha='center', va='center')
                
                        ax_3d.set_xlabel('Distance to Ideal Best (Lower is Better)')
                        ax_3d.set_ylabel('Distance to Ideal Worst (Higher is Better)')
                        ax_3d.set_zlabel('Score (Higher is Better)')
                
                        import matplotlib.patches as mpatches
                        from matplotlib.legend_handler import HandlerPatch
                        
                        class HandlerCircleWithText(HandlerPatch):
                            def __init__(self, rank, color, **kwargs):
                                super().__init__(**kwargs)
                                self.rank = rank
                                self.color = color
                        
                            def create_artists(self, legend, orig_handle, xdescent, ydescent, width, height, fontsize, trans):
                                center = (width / 2 - xdescent, height / 2 - ydescent)
                                radius = height / 1.5
                                circle = mpatches.Circle(center, radius=radius, facecolor=self.color, lw=0.5, transform=trans)
                                txt = plt.Text(x=center[0], y=center[1], text=str(self.rank),
                                               ha="center", va="center", fontsize=7, color='white', transform=trans)
                                return [circle, txt]
                        
                        cmap = plt.get_cmap('viridis')
                        norm = plt.Normalize(df_processed['Rank'].min(), df_processed['Rank'].max())
                        legend_elements = []
                        handler_map = {}
                        
                        for method_name in df_processed.index:
                            rank = int(df_processed.loc[method_name, 'Rank'])
                            color = cmap(norm(rank))
                        
                            dummy_circle = mpatches.Circle((0, 0), radius=6)
                            legend_elements.append((dummy_circle, method_name, rank)) 
                            handler_map[dummy_circle] = HandlerCircleWithText(rank=rank, color=color)
                        
                        legend_elements = sorted(legend_elements, key=lambda x: x[2])
                        handles, labels, _ = zip(*legend_elements)

                        ax_3d.legend(handles, labels, title="EVs",
                                     handler_map=handler_map, loc='center left', bbox_to_anchor=(1.05, 0.5), frameon=True)


                
                        plt.title('3D Plot: Distances & Scores')
                        st.pyplot(fig_3d)

                doc = Document()
                doc.add_heading('EVs Comparison and Ranking Report', 0)
                doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph(f'Number of EVs compared: {len(df_processed)}')
                doc.add_heading('Parameters Classification', level=1)
                doc.add_paragraph('Parameters Maximized: ' + (', '.join(maximize_factors_selected) if maximize_factors_selected else "None"))
                doc.add_paragraph('Parameters Minimized: ' + (', '.join(minimize_factors_selected) if minimize_factors_selected else "None"))
                
                doc.add_heading('Parameters Weights', level=1)
                if weights:
                    weights_table = doc.add_table(rows=1, cols=2)
                    weights_table.style = 'Table Grid'
                    hdr_cells = weights_table.rows[0].cells; hdr_cells[0].text = 'Parameter'; hdr_cells[1].text = 'Weight'
                    for factor, weight in weights.items():
                        row_cells = weights_table.add_row().cells; row_cells[0].text = factor; row_cells[1].text = f"{weight:.2f}"
                
                doc.add_heading('EVs Rankings', level=1)
                rank_table = doc.add_table(rows=1, cols=3); rank_table.style = 'Table Grid'
                hdr_cells = rank_table.rows[0].cells; hdr_cells[0].text = 'EVs'; hdr_cells[1].text = 'Score'; hdr_cells[2].text = 'Rank'
                for idx, row in ranking_df.iterrows():
                    row_cells = rank_table.add_row().cells; row_cells[0].text = str(idx); row_cells[1].text = f"{row['Score']:.4f}"; row_cells[2].text = f"{int(row['Rank'])}"

                doc.add_heading('Input Data', level=1)
                if not st.session_state.data_df.empty:
                    data_table = doc.add_table(rows=1, cols=len(st.session_state.data_df.columns) + 1); data_table.style = 'Table Grid'
                    hdr_cells = data_table.rows[0].cells; hdr_cells[0].text = 'EVs'
                    for i, factor_name_doc in enumerate(st.session_state.data_df.columns): hdr_cells[i+1].text = factor_name_doc
                    for idx, row_data in st.session_state.data_df.iterrows():
                        row_cells = data_table.add_row().cells; row_cells[0].text = str(idx)
                        for i, val in enumerate(row_data): row_cells[i+1].text = f"{val}"
                
                doc.add_heading('Visualizations', level=1)
                def add_figure_to_doc(fig, title):
                    if fig:
                        buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=300, bbox_inches='tight'); buf.seek(0)
                        doc.add_heading(title, level=2); doc.add_picture(buf, width=Inches(6.0)); doc.add_paragraph()

                if 'fig_bar' in locals(): add_figure_to_doc(fig_bar, 'Bar Plot of Scores')
                if 'fig_heat' in locals() and normalized_factors_present and not df_processed[normalized_factors_present].empty : add_figure_to_doc(fig_heat, 'Heatmap of Normalized Parameters')
                if 'fig_radar' in locals() and not top_methods_df.empty and not top_methods_df[normalized_factors_present].empty: add_figure_to_doc(fig_radar, 'Top EVs Radar Chart')
                if 'fig_3d' in locals() and len(normalized_factors_present) >= 2: add_figure_to_doc(fig_3d, '3D Scatter Plot')

                docx_buffer = io.BytesIO(); doc.save(docx_buffer); docx_buffer.seek(0)
                st.subheader("Download Report")
                b64_docx = base64.b64encode(docx_buffer.getvalue()).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="evs_comparison_report.docx">Download Report (DOCX)</a>'
                st.markdown(href, unsafe_allow_html=True)
                st.markdown("Note: To save as PDF, download the DOCX file and open it in your preferred word processor, then save as PDF.")
else:
    if input_method == "Manual Entry":
        st.info("⬅️ Please define EVs names, Parameter names, classify Parameters, enter data, and click 'Submit Data' in the sidebar to begin.")
    elif input_method == "Upload Excel File":
         st.info("⬅️ Please upload an Excel file using the sidebar to begin.")
    if st.session_state.data_df is not None and st.session_state.data_df.empty :
         st.warning("The current data is empty. Please check your input or upload a valid file.")
